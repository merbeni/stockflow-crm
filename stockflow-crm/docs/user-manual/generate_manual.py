"""
Genera el Manual de Usuario de StockFlow CRM en formato .docx.
Ejecutar: python generate_manual.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── colores de marca ──────────────────────────────────────────────────────────
COLOR_PRIMARIO   = RGBColor(0x1A, 0x56, 0xDB)   # azul
COLOR_SECUNDARIO = RGBColor(0x0E, 0x9F, 0x6E)   # verde
COLOR_GRIS       = RGBColor(0x6B, 0x72, 0x80)   # gris texto
COLOR_ENCABEZADO = RGBColor(0x11, 0x18, 0x27)   # casi negro
COLOR_FONDO_TAB  = RGBColor(0xE8, 0xF0, 0xFE)   # azul claro


# ── helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, color_hex: str):
    """Pone color de fondo a una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.runs[0] if h.runs else h.add_run(text)
    if color:
        run.font.color.rgb = color
    return h


def add_paragraph(doc, text, bold=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def add_step(doc, number, title, description):
    p = doc.add_paragraph(style='List Number')
    run_title = p.add_run(f"{title}: ")
    run_title.bold = True
    run_title.font.color.rgb = COLOR_PRIMARIO
    p.add_run(description)


def add_note(doc, text, tipo="Nota"):
    p = doc.add_paragraph()
    run_label = p.add_run(f"⚑ {tipo}: ")
    run_label.bold = True
    run_label.font.color.rgb = COLOR_SECUNDARIO
    run_label.font.size = Pt(10)
    run_desc = p.add_run(text)
    run_desc.font.size = Pt(10)
    run_desc.font.color.rgb = COLOR_GRIS


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Encabezado
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, '1A56DB')
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

    # Filas de datos
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        bg = 'E8F0FE' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


# ── documento ─────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # ── Márgenes ────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    # ── Portada ─────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = titulo.add_run("StockFlow CRM")
    r.font.size = Pt(32)
    r.font.bold = True
    r.font.color.rgb = COLOR_PRIMARIO

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = subtitulo.add_run("Manual de Usuario")
    r2.font.size = Pt(20)
    r2.font.color.rgb = COLOR_ENCABEZADO

    doc.add_paragraph()
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = desc.add_run("Mini CRM para negocios de ecommerce\nGestión de inventario, proveedores, clientes, pedidos\ny procesamiento de facturas con Inteligencia Artificial")
    r3.font.size = Pt(12)
    r3.font.color.rgb = COLOR_GRIS

    doc.add_paragraph()
    doc.add_paragraph()
    fecha = doc.add_paragraph()
    fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecha.add_run(f"Versión 1.0  ·  {datetime.datetime.now().strftime('%B %Y')}").font.color.rgb = COLOR_GRIS

    doc.add_paragraph()
    proyecto = doc.add_paragraph()
    proyecto.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rp = proyecto.add_run("Proyecto Final de Grado — Técnico en Programación")
    rp.bold = True
    rp.font.size = Pt(11)

    doc.add_page_break()

    # ── Índice ───────────────────────────────────────────────────────────────
    add_heading(doc, "Índice de contenidos", level=1, color=COLOR_PRIMARIO)
    secciones = [
        ("1.", "Introducción"),
        ("2.", "Acceso al sistema"),
        ("3.", "Interfaz general"),
        ("4.", "Módulo: Inventario (Productos)"),
        ("5.", "Módulo: Proveedores"),
        ("6.", "Módulo: Facturas — Procesamiento con IA"),
        ("7.", "Módulo: Movimientos de Stock"),
        ("8.", "Módulo: Clientes"),
        ("9.", "Módulo: Pedidos"),
        ("10.", "Flujos de trabajo principales"),
        ("11.", "Preguntas frecuentes"),
    ]
    for num, titulo_sec in secciones:
        p = doc.add_paragraph(style='List Bullet')
        r_num = p.add_run(f"{num}  ")
        r_num.bold = True
        r_num.font.color.rgb = COLOR_PRIMARIO
        p.add_run(titulo_sec)

    doc.add_page_break()

    # ── 1. Introducción ──────────────────────────────────────────────────────
    add_heading(doc, "1. Introducción", level=1, color=COLOR_PRIMARIO)
    add_paragraph(doc,
        "StockFlow CRM es una aplicación web diseñada para ayudar a negocios de ecommerce "
        "a gestionar su operación diaria de forma integrada: inventario de productos, "
        "relaciones con proveedores, procesamiento de facturas de compra mediante "
        "Inteligencia Artificial, seguimiento de clientes y control de pedidos.")
    doc.add_paragraph()
    add_paragraph(doc, "Características principales:", bold=True)
    features = [
        "Procesamiento automático de facturas: subí una imagen o PDF y el sistema extrae todos los ítems usando Gemini 2.5 Flash (IA de Google).",
        "Control de stock en tiempo real: cada factura confirmada y cada pedido procesado actualiza el stock automáticamente.",
        "Alertas de stock mínimo: el sistema avisa cuando un producto cae por debajo del umbral configurado.",
        "Historial completo: todos los movimientos de stock quedan registrados y son auditables.",
        "Notificaciones por email: el cliente recibe un correo en cada cambio de estado de su pedido.",
        "Interfaz intuitiva: diseñada para que cualquier persona pueda operar el sistema sin capacitación técnica.",
    ]
    for f in features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f)

    doc.add_paragraph()
    add_heading(doc, "Roles de usuario", level=2)
    add_table(doc,
        ["Rol", "Descripción"],
        [
            ["admin", "Acceso completo a todas las funciones del sistema."],
            ["operator", "Acceso operativo: puede gestionar inventario, facturas, clientes y pedidos."],
        ],
        col_widths=[4, 12]
    )

    doc.add_page_break()

    # ── 2. Acceso al sistema ─────────────────────────────────────────────────
    add_heading(doc, "2. Acceso al sistema", level=1, color=COLOR_PRIMARIO)

    add_heading(doc, "2.1  Iniciar sesión", level=2)
    add_paragraph(doc,
        "Para acceder a StockFlow CRM, abrí el navegador e ingresá a la dirección "
        "que te haya proporcionado el administrador (por ejemplo: http://localhost:5173 "
        "en entorno local, o la URL de producción en Azure).")
    doc.add_paragraph()
    add_paragraph(doc, "Pasos para iniciar sesión:", bold=True)
    add_step(doc, 1, "Ingresá tu email", "escribí la dirección de correo con la que fuiste registrado.")
    add_step(doc, 2, "Ingresá tu contraseña", "la contraseña es sensible a mayúsculas y minúsculas.")
    add_step(doc, 3, "Hacé clic en «Iniciar sesión»", "si los datos son correctos, serás redirigido al panel principal.")
    doc.add_paragraph()
    add_note(doc, "Si olvidaste tu contraseña, contactá al administrador del sistema para que la restablezca.")

    add_heading(doc, "2.2  Cerrar sesión", level=2)
    add_paragraph(doc,
        "Hacé clic en el botón «Cerrar sesión» ubicado en la parte inferior izquierda "
        "de la barra lateral. Tu sesión se cerrará y serás redirigido a la pantalla de login.")
    add_note(doc,
        "La sesión expira automáticamente después de 60 minutos de inactividad. "
        "Guardá tu trabajo antes de alejarte del equipo.", tipo="Importante")

    doc.add_page_break()

    # ── 3. Interfaz general ──────────────────────────────────────────────────
    add_heading(doc, "3. Interfaz general", level=1, color=COLOR_PRIMARIO)
    add_paragraph(doc,
        "Una vez que iniciás sesión, encontrarás la siguiente estructura en pantalla:")
    doc.add_paragraph()
    add_table(doc,
        ["Zona", "Descripción"],
        [
            ["Barra lateral (izquierda)", "Menú de navegación con acceso a todos los módulos. Muestra el email del usuario logueado y el botón de cerrar sesión."],
            ["Área de contenido (derecha)", "Aquí se muestra la sección activa. Cambia al hacer clic en los ítems del menú."],
        ],
        col_widths=[5, 11]
    )
    add_paragraph(doc, "Módulos disponibles en el menú:", bold=True)
    modulos = [
        ("Productos", "Gestión del catálogo de productos e inventario."),
        ("Proveedores", "Administración de proveedores."),
        ("Facturas", "Carga y procesamiento de facturas con IA."),
        ("Movimientos", "Historial de movimientos de stock."),
        ("Clientes", "Administración de clientes."),
        ("Pedidos", "Gestión y seguimiento de pedidos."),
    ]
    add_table(doc,
        ["Módulo", "Función"],
        modulos,
        col_widths=[4, 12]
    )

    doc.add_page_break()

    # ── 4. Inventario ────────────────────────────────────────────────────────
    add_heading(doc, "4. Módulo: Inventario (Productos)", level=1, color=COLOR_PRIMARIO)
    add_paragraph(doc,
        "Este módulo te permite gestionar todos los productos de tu catálogo, "
        "incluyendo su precio, stock disponible y estado activo/inactivo.")

    add_heading(doc, "4.1  Ver la lista de productos", level=2)
    add_paragraph(doc,
        "Al entrar al módulo «Productos» verás una tabla con todos los productos registrados. "
        "Cada fila muestra:")
    campos = [
        ["SKU", "Código único del producto."],
        ["Nombre", "Nombre descriptivo del producto."],
        ["Precio", "Precio unitario."],
        ["Stock actual", "Cantidad disponible en este momento."],
        ["Stock mínimo", "Umbral de alerta. Si el stock actual es menor, se muestra el indicador rojo «Stock bajo»."],
        ["Estado", "Activo (disponible para pedidos) o Inactivo (no se puede agregar a pedidos)."],
    ]
    add_table(doc, ["Campo", "Descripción"], campos, col_widths=[4, 12])
    add_note(doc, "Los productos con stock bajo aparecen marcados con una etiqueta roja. "
             "Usá el filtro «Solo stock bajo» para verlos todos juntos.")

    add_heading(doc, "4.2  Crear un producto", level=2)
    add_step(doc, 1, "Hacé clic en «Nuevo producto»", "se abre el formulario de creación.")
    add_step(doc, 2, "Completá el SKU", "es el código único del producto (ej. ZAP-001). No puede repetirse.")
    add_step(doc, 3, "Completá el nombre, descripción y precio", "el precio debe ser un número mayor o igual a 0.")
    add_step(doc, 4, "Indicá el stock inicial", "si ya tenés unidades en depósito, ingresá la cantidad. El sistema registrará un movimiento de entrada automáticamente.")
    add_step(doc, 5, "Configurá el stock mínimo", "cuando el stock actual caiga por debajo de este valor recibirás una alerta.")
    add_step(doc, 6, "Hacé clic en «Guardar»", "el producto queda registrado en el sistema.")

    add_heading(doc, "4.3  Editar un producto", level=2)
    add_paragraph(doc, "Hacé clic en el ícono de edición (lápiz) en la fila del producto.")
    add_paragraph(doc, "Podés modificar nombre, descripción, precio, stock actual, stock mínimo y estado activo.")
    add_note(doc, "Si cambiás el stock actual, el sistema registra automáticamente un movimiento de ajuste con la diferencia (positiva o negativa).")
    add_note(doc, "El SKU no puede modificarse una vez creado el producto.", tipo="Importante")

    add_heading(doc, "4.4  Eliminar un producto", level=2)
    add_paragraph(doc, "Hacé clic en el ícono de eliminación (papelera) y confirmá la acción.")
    add_note(doc, "Un producto no puede eliminarse si tiene movimientos de stock o historial de pedidos. En ese caso, desactivalo usando el campo «Estado».", tipo="Importante")

    doc.add_page_break()

    # ── 5. Proveedores ───────────────────────────────────────────────────────
    add_heading(doc, "5. Módulo: Proveedores", level=1, color=COLOR_PRIMARIO)
    add_paragraph(doc,
        "Registrá a tus proveedores para poder asociarlos a las facturas y que el sistema "
        "aprenda automáticamente la relación entre sus códigos (SKUs propios) y tus productos.")

    add_heading(doc, "5.1  Crear un proveedor", level=2)
    add_step(doc, 1, "Hacé clic en «Nuevo proveedor»", "se abre el formulario.")
    add_step(doc, 2, "Completá los datos", "Nombre (obligatorio), nombre de contacto (obligatorio), email (obligatorio) y teléfono (opcional).")
    add_step(doc, 3, "Hacé clic en «Guardar»", "el proveedor queda registrado.")

    add_heading(doc, "5.2  Editar y eliminar proveedores", level=2)
    add_paragraph(doc,
        "Usá los íconos de edición y eliminación en la tabla. "
        "Un proveedor puede eliminarse siempre que no tenga facturas asociadas.")

    doc.add_page_break()

    # ── 6. Facturas ──────────────────────────────────────────────────────────
    add_heading(doc, "6. Módulo: Facturas — Procesamiento con IA", level=1, color=COLOR_PRIMARIO)
    add_paragraph(doc,
        "Este es el módulo central de StockFlow CRM. Permite cargar facturas de proveedores "
        "en formato PDF, JPG o PNG y procesarlas automáticamente con Inteligencia Artificial "
        "(Gemini 2.5 Flash de Google). El sistema extrae todos los ítems de la factura, "
        "los muestra para revisión y, al confirmar, actualiza el stock.")

    add_heading(doc, "6.1  Flujo completo de procesamiento", level=2)
    add_paragraph(doc, "El proceso tiene dos pasos principales:", bold=True)

    add_heading(doc, "Paso 1 — Subir el archivo", level=3)
    add_step(doc, 1, "Hacé clic en «Procesar factura»", "se abre el panel de carga.")
    add_step(doc, 2, "Seleccioná el archivo", "hacé clic en el área de carga o arrastrá el archivo. Formatos admitidos: PDF, JPG, PNG. Tamaño máximo: 20 MB.")
    add_step(doc, 3, "Revisá la previsualización", "el sistema muestra la imagen o el PDF para que verifiques que es el archivo correcto.")
    add_step(doc, 4, "Hacé clic en «Analizar con IA»", "el sistema envía el archivo a Gemini 2.5 Flash y espera la respuesta (puede tomar 5 a 20 segundos).")
    add_note(doc,
        "Si el análisis falla con un mensaje de error o demora demasiado, es posible que el servicio de Gemini "
        "esté saturado. Esto es completamente normal cuando se usa el plan gratuito (free trial) de Google AI Studio, "
        "que tiene un límite de solicitudes por minuto. Esperá unos segundos y volvé a intentarlo.",
        tipo="Importante")
    add_note(doc, "Si el error persiste por varios minutos, verificá que la clave GOOGLE_API_KEY en el archivo "
             ".env sea correcta y que no hayas agotado la cuota diaria gratuita. Podés revisar el estado y "
             "los límites de tu clave en https://aistudio.google.com")

    add_heading(doc, "Paso 2 — Revisar y confirmar", level=3)
    add_paragraph(doc, "Tras el análisis, verás la tabla de ítems detectados. Para cada ítem:")
    items_tabla = [
        ["Descripción", "Texto extraído de la factura."],
        ["Cantidad", "Unidades detectadas por la IA."],
        ["Precio unitario", "Precio extraído de la factura."],
        ["Confianza", "Nivel de certeza de la IA: Alta (verde), Media (amarillo) o Baja (rojo). Los ítems de baja confianza requieren revisión manual."],
        ["Producto", "Sugerencia automática de qué producto del catálogo corresponde a este ítem."],
        ["SKU proveedor", "Código del proveedor para ese ítem (se guarda para futuros reconocimientos)."],
        ["Omitir", "Casilla para ignorar un ítem y no actualizar su stock."],
    ]
    add_table(doc, ["Campo", "Descripción"], items_tabla, col_widths=[4, 12])

    add_step(doc, 1, "Asigná el proveedor", "seleccioná un proveedor existente o completá los datos de uno nuevo.")
    add_step(doc, 2, "Revisá cada ítem", "para cada ítem, seleccioná el producto de tu catálogo que corresponde, o creá uno nuevo si no existe.")
    add_step(doc, 3, "Completá el SKU del proveedor", "si sabés el código que usa el proveedor para ese producto, ingresalo. El sistema lo recordará para futuras facturas del mismo proveedor.")
    add_step(doc, 4, "Omití ítems irrelevantes", "si un ítem es un cargo de flete, descuento u otro concepto que no actualiza stock, tildá la casilla «Omitir».")
    add_step(doc, 5, "Hacé clic en «Confirmar factura»", "el stock de cada producto se actualiza automáticamente. Se generan movimientos de tipo «Entrada» por cada ítem confirmado.")
    doc.add_paragraph()
    add_note(doc, "Si el stock de algún producto queda por debajo del mínimo luego de la confirmación, recibirás una alerta de stock bajo.")
    add_note(doc, "Si la factura tiene errores graves, podés rechazarla con el botón «Rechazar». Esto cancela el proceso sin afectar el stock.", tipo="Alternativa")

    add_heading(doc, "6.2  Ver el historial de facturas", level=2)
    add_paragraph(doc,
        "En la vista principal del módulo «Facturas» verás todas las facturas procesadas "
        "con su estado (Pendiente, Confirmada o Rechazada), la fecha y el proveedor asociado.")

    doc.add_page_break()

    # ── 7. Movimientos de stock ──────────────────────────────────────────────
    add_heading(doc, "7. Módulo: Movimientos de Stock", level=1, color=COLOR_PRIMARIO)
    add_paragraph(doc,
        "Este módulo es de solo lectura y muestra el historial completo de todos los "
        "cambios en el stock de los productos. Es útil para auditorías y para "
        "entender por qué cambió el inventario en un momento dado.")

    add_heading(doc, "7.1  Tipos de movimiento", level=2)
    add_table(doc,
        ["Tipo", "Color", "Cuándo se genera"],
        [
            ["Entrada", "Verde", "Cuando se confirma una factura de proveedor. El stock aumenta."],
            ["Venta", "Rojo", "Cuando un pedido pasa al estado «Procesando». El stock disminuye."],
            ["Ajuste", "Gris", "Cuando se edita manualmente el stock de un producto desde el módulo de Inventario. Puede ser positivo o negativo."],
        ],
        col_widths=[3, 2.5, 10.5]
    )

    add_heading(doc, "7.2  Filtros disponibles", level=2)
    add_table(doc,
        ["Filtro", "Descripción"],
        [
            ["Tipo", "Mostrar solo entradas, ventas o ajustes."],
            ["Producto", "Ver solo los movimientos de un producto específico."],
            ["Fecha desde / hasta", "Acotar el historial a un rango de fechas."],
        ],
        col_widths=[4, 12]
    )
    add_paragraph(doc,
        "Hacé clic en cualquier fila para ver el detalle del movimiento, "
        "incluyendo a qué factura o pedido está vinculado.")

    doc.add_page_break()

    # ── 8. Clientes ──────────────────────────────────────────────────────────
    add_heading(doc, "8. Módulo: Clientes", level=1, color=COLOR_PRIMARIO)
    add_paragraph(doc,
        "Registrá y administrá los clientes de tu negocio. Podés ver el historial "
        "completo de pedidos de cada cliente.")

    add_heading(doc, "8.1  Datos de un cliente", level=2)
    add_table(doc,
        ["Campo", "Obligatorio", "Descripción"],
        [
            ["Nombre", "Sí", "Nombre completo o razón social del cliente."],
            ["Email", "Sí", "Se usa para enviar notificaciones de estado de pedidos."],
            ["Teléfono", "Sí", "Número de contacto."],
            ["Dirección", "No", "Domicilio de entrega."],
        ],
        col_widths=[3, 3, 10]
    )

    add_heading(doc, "8.2  Ver el historial de pedidos de un cliente", level=2)
    add_step(doc, 1, "Hacé clic en el ícono de historial", "en la fila del cliente.")
    add_step(doc, 2, "Se abre el panel de historial", "con todos los pedidos del cliente, su estado, fecha y total.")

    doc.add_page_break()

    # ── 9. Pedidos ───────────────────────────────────────────────────────────
    add_heading(doc, "9. Módulo: Pedidos", level=1, color=COLOR_PRIMARIO)
    add_paragraph(doc,
        "Gestioná los pedidos de tus clientes desde la creación hasta la entrega. "
        "Cada pedido sigue un flujo de estados con control de stock integrado.")

    add_heading(doc, "9.1  Estados de un pedido", level=2)
    add_table(doc,
        ["Estado", "Color", "Descripción"],
        [
            ["Pendiente", "Amarillo", "El pedido fue creado. Se pueden agregar o quitar ítems."],
            ["Procesando", "Azul", "El pedido fue confirmado. El stock se descuenta automáticamente en este momento."],
            ["Enviado", "Verde oscuro", "El pedido fue despachado."],
            ["Entregado", "Verde claro", "El cliente recibió el pedido. Estado final."],
        ],
        col_widths=[3, 3, 10]
    )

    add_heading(doc, "9.2  Crear un pedido", level=2)
    add_step(doc, 1, "Hacé clic en «Nuevo pedido»", "se abre el formulario.")
    add_step(doc, 2, "Seleccioná el cliente", "buscá al cliente en el listado desplegable.")
    add_step(doc, 3, "Hacé clic en «Crear»", "el pedido se crea en estado «Pendiente» y queda vacío.")

    add_heading(doc, "9.3  Agregar ítems al pedido", level=2)
    add_step(doc, 1, "En la fila del pedido, hacé clic en «Agregar ítem»", "se abre el formulario de ítem.")
    add_step(doc, 2, "Seleccioná el producto", "el sistema muestra el stock disponible en tiempo real.")
    add_step(doc, 3, "Ingresá la cantidad", "no puede superar el stock disponible.")
    add_step(doc, 4, "Ingresá el precio unitario", "por defecto se sugiere el precio del catálogo, pero puede modificarse.")
    add_step(doc, 5, "Hacé clic en «Agregar»", "el ítem queda en el pedido. El total se actualiza automáticamente.")
    add_note(doc, "Podés agregar varios ítems al mismo pedido. Usá el ícono de papelera para quitar un ítem antes de procesar el pedido.")

    add_heading(doc, "9.4  Avanzar el estado del pedido", level=2)
    add_paragraph(doc,
        "Hacé clic en el botón «Avanzar» para mover el pedido al siguiente estado. "
        "Cada clic avanza un paso:")
    add_table(doc,
        ["De", "A", "Qué ocurre"],
        [
            ["Pendiente", "Procesando", "Se valida el stock. Si hay suficiente, se descuenta de cada producto y se registran movimientos de venta. Se envía email al cliente."],
            ["Procesando", "Enviado", "Sin cambios de stock. Se envía email al cliente."],
            ["Enviado", "Entregado", "Estado final. Se envía email al cliente."],
        ],
        col_widths=[3, 3, 10]
    )
    add_note(doc, "Solo se puede eliminar un pedido si está en estado «Pendiente».", tipo="Importante")
    add_note(doc, "Si el stock de algún producto no alcanza al momento de confirmar (Pendiente → Procesando), el sistema muestra un error y no avanza el pedido.")

    doc.add_page_break()

    # ── 10. Flujos de trabajo ────────────────────────────────────────────────
    add_heading(doc, "10. Flujos de trabajo principales", level=1, color=COLOR_PRIMARIO)

    add_heading(doc, "Flujo A — Recepción de mercadería con factura", level=2)
    pasos_a = [
        ("Recibís mercadería del proveedor junto con una factura en papel o digital", ""),
        ("Vas al módulo «Facturas»", "y hacés clic en «Procesar factura»."),
        ("Subís la imagen o PDF de la factura", "y esperás el análisis de la IA."),
        ("Revisás los ítems detectados", "corrigiendo o confirmando cada uno y asignando los productos del catálogo."),
        ("Confirmás la factura", "el stock se actualiza automáticamente para cada producto confirmado."),
        ("Verificás en el módulo «Productos»", "que los stocks quedaron correctos."),
        ("Consultás en «Movimientos»", "para ver el registro de entradas generado."),
    ]
    for i, (titulo, desc) in enumerate(pasos_a, 1):
        add_step(doc, i, titulo, f" {desc}" if desc else "")

    doc.add_paragraph()
    add_heading(doc, "Flujo B — Gestión de un pedido de cliente", level=2)
    pasos_b = [
        ("El cliente realiza un pedido", "vas al módulo «Pedidos» y hacés clic en «Nuevo pedido»."),
        ("Seleccionás al cliente", "y creás el pedido."),
        ("Agregás los productos", "verificando que el stock disponible sea suficiente."),
        ("Al preparar el pedido", "hacés clic en «Avanzar» → pasa a «Procesando» y el stock se descuenta."),
        ("Al despachar el pedido", "hacés clic en «Avanzar» → pasa a «Enviado». El cliente recibe un email."),
        ("Al confirmar la entrega", "hacés clic en «Avanzar» → pasa a «Entregado». El cliente recibe un email de confirmación."),
    ]
    for i, (titulo, desc) in enumerate(pasos_b, 1):
        add_step(doc, i, titulo, f" {desc}" if desc else "")

    doc.add_paragraph()
    add_heading(doc, "Flujo C — Alta de producto nuevo desde una factura", level=2)
    add_paragraph(doc,
        "Cuando la IA detecta un ítem en la factura que no existe en el catálogo, "
        "podés crearlo directamente desde la pantalla de revisión:")
    pasos_c = [
        ("En la columna «Producto»", "dejá el selector vacío y tildá «Crear producto nuevo»."),
        ("Completá el SKU, nombre, precio y stock mínimo", "del nuevo producto."),
        ("Al confirmar la factura", "el producto se crea en el catálogo y el stock se inicializa con la cantidad de la factura."),
    ]
    for i, (titulo, desc) in enumerate(pasos_c, 1):
        add_step(doc, i, titulo, f" {desc}" if desc else "")

    doc.add_page_break()

    # ── 11. Preguntas frecuentes ─────────────────────────────────────────────
    add_heading(doc, "11. Preguntas frecuentes", level=1, color=COLOR_PRIMARIO)

    faqs = [
        (
            "¿Qué hago si la IA extrajo datos incorrectos de la factura?",
            "Podés editar manualmente cualquier campo antes de confirmar. Si el error es en "
            "la cantidad o precio de un ítem, simplemente modificá el valor en la celda correspondiente. "
            "Si el ítem está completamente equivocado, tildá «Omitir» para ese ítem."
        ),
        (
            "¿Por qué un ítem aparece marcado en rojo en la revisión de facturas?",
            "El color rojo indica que la IA tiene baja confianza en ese dato (nivel «Baja»). "
            "Significa que el valor fue estimado porque no se pudo leer claramente. "
            "Revisalo manualmente contra la factura física antes de confirmar."
        ),
        (
            "¿El sistema guarda las facturas originales?",
            "Actualmente el sistema guarda los datos extraídos por la IA (en formato JSON interno), "
            "pero no almacena el archivo original. Se recomienda guardar los archivos originales "
            "en una carpeta organizada en tu computadora o nube."
        ),
        (
            "¿Puedo deshacer una factura ya confirmada?",
            "No. Una vez confirmada, la factura no puede revertirse automáticamente. "
            "Si cometiste un error, deberás hacer un ajuste manual en el stock del producto afectado "
            "desde el módulo «Productos»."
        ),
        (
            "¿Por qué no puedo eliminar un producto?",
            "Un producto con movimientos de stock o historial de pedidos no puede eliminarse "
            "para preservar la integridad del historial. En su lugar, podés desactivarlo "
            "desmarcando la opción «Activo» en la edición del producto. "
            "Un producto inactivo no puede agregarse a nuevos pedidos."
        ),
        (
            "¿Los emails de notificación son obligatorios?",
            "No. El sistema de emails es opcional. Si no se configura la clave de SendGrid, "
            "los emails se omiten silenciosamente sin afectar ninguna otra funcionalidad."
        ),
        (
            "La IA devolvió un error al procesar la factura. ¿Qué hago?",
            "Si al hacer clic en «Analizar con IA» aparece un mensaje de error, lo más probable es que el "
            "servicio de Gemini esté saturado. Esto es normal cuando se usa el plan gratuito (free trial) de "
            "Google AI Studio: el plan tiene un límite de solicitudes por minuto y por día. "
            "Esperá entre 10 y 30 segundos y volvé a intentar la carga. Si el error persiste, "
            "puede que hayas alcanzado la cuota diaria; en ese caso, intentalo nuevamente al día siguiente "
            "o utilizá una clave de API diferente."
        ),
        (
            "¿Puedo tener varios usuarios en el sistema?",
            "Sí. El administrador puede registrar usuarios adicionales a través del endpoint "
            "/auth/register (o directamente desde la API). Cada usuario tiene un rol "
            "(admin u operator) que define sus permisos."
        ),
    ]
    for pregunta, respuesta in faqs:
        p = doc.add_paragraph()
        r_q = p.add_run(f"❓ {pregunta}")
        r_q.bold = True
        r_q.font.color.rgb = COLOR_PRIMARIO
        r_q.font.size = Pt(11)

        p_r = doc.add_paragraph()
        r_a = p_r.add_run(respuesta)
        r_a.font.size = Pt(10)
        r_a.font.color.rgb = COLOR_ENCABEZADO
        p_r.paragraph_format.left_indent = Cm(0.5)
        doc.add_paragraph()

    # ── Pie de página ────────────────────────────────────────────────────────
    doc.add_page_break()
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_pie = pie.add_run(
        f"StockFlow CRM · Manual de Usuario v1.0 · "
        f"{datetime.datetime.now().strftime('%Y')}\n"
        "Proyecto Final de Grado — Técnico en Programación"
    )
    r_pie.font.size = Pt(9)
    r_pie.font.color.rgb = COLOR_GRIS

    # ── Guardar ──────────────────────────────────────────────────────────────
    output_path = "Manual_de_Usuario_StockFlow_CRM.docx"
    doc.save(output_path)
    print(f"Documento generado: {output_path}")


if __name__ == "__main__":
    build()
